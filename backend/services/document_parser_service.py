"""
Document Parser Service
=======================

Parses various document formats to extract curriculum content:
- PDF files (using pdfplumber for layout-aware extraction)
- Word documents (using python-docx)
- Plain text (structure detection via patterns)

Returns structured content suitable for AI processing in the curriculum upload pipeline.
"""

import re
from io import BytesIO
from typing import Dict, List, Optional, Any
from services.base_service import BaseService

from utils.logger import get_logger

logger = get_logger(__name__)


class DocumentParserError(Exception):
    """Base exception for document parsing errors."""
    pass


class UnsupportedFormatError(DocumentParserError):
    """File format is not supported."""
    pass


class DocumentParserService(BaseService):
    """
    Service for parsing documents into structured curriculum content.

    Supports:
    - PDF files (.pdf)
    - Word documents (.docx)
    - Plain text

    Output format is consistent across all types to feed into AI processing.
    """

    # Patterns for detecting structure in text
    HEADING_PATTERNS = [
        r'^#{1,6}\s+(.+)$',  # Markdown headings
        r'^([A-Z][A-Z\s]+)$',  # ALL CAPS lines
        r'^(\d+\.)\s+(.+)$',  # Numbered sections (1. Title)
        r'^(Module|Unit|Chapter|Lesson|Week|Section)\s*(\d+)?[:\s]*(.+)?$',  # Common curriculum headers
    ]

    # Patterns for detecting lesson/assignment boundaries
    LESSON_PATTERNS = [
        r'(?i)^(lesson|activity|assignment|task|exercise|lab|project|discussion|quiz|test|exam)\s*(\d+)?[:\s]*(.+)?$',
    ]

    def parse_document(
        self,
        content: bytes,
        source_type: str,
        filename: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Parse document content based on source type.

        Args:
            content: Raw bytes of the document
            source_type: Type of document ('pdf', 'docx', 'text')
            filename: Original filename (optional, for metadata)

        Returns:
            Dict with structured content:
            {
                'success': bool,
                'error': str (if failed),
                'raw_text': str,
                'sections': [...],
                'metadata': {...}
            }
        """
        try:
            if source_type == 'pdf':
                return self.parse_pdf(content, filename)
            elif source_type == 'docx':
                return self.parse_docx(content, filename)
            elif source_type == 'text':
                # For text, content should be string, but handle bytes too
                if isinstance(content, bytes):
                    text = content.decode('utf-8', errors='replace')
                else:
                    text = content
                return self.parse_text(text, filename)
            else:
                raise UnsupportedFormatError(f"Unsupported document type: {source_type}")

        except DocumentParserError:
            raise
        except Exception as e:
            logger.error(f"Error parsing document: {str(e)}")
            return {
                'success': False,
                'error': f'Failed to parse document: {str(e)}',
                'raw_text': '',
                'sections': [],
                'metadata': {}
            }

    def parse_pdf(self, file_content: bytes, filename: Optional[str] = None) -> Dict[str, Any]:
        """
        Parse PDF document using pdfplumber.

        Extracts:
        - All text content with layout awareness
        - Page numbers
        - Section headings (detected from formatting)
        - Tables (converted to text)

        Memory-optimized for large PDFs:
        - Processes pages incrementally
        - Clears page data after extraction
        - Explicitly releases BytesIO buffer

        Args:
            file_content: Raw bytes of the PDF file
            filename: Original filename

        Returns:
            Structured content dict
        """
        import gc

        try:
            import pdfplumber
        except ImportError:
            logger.error("pdfplumber not installed. Install with: pip install pdfplumber")
            return {
                'success': False,
                'error': 'PDF parsing library not available',
                'raw_text': '',
                'sections': [],
                'metadata': {}
            }

        # Track BytesIO explicitly for cleanup
        pdf_buffer = None

        try:
            sections = []
            current_section = None
            page_count = 0

            # Use StringIO-like approach: build text incrementally to avoid
            # holding multiple copies (list + joined string)
            text_parts = []
            total_chars = 0

            # Create BytesIO buffer explicitly so we can close it
            pdf_buffer = BytesIO(file_content)

            with pdfplumber.open(pdf_buffer) as pdf:
                page_count = len(pdf.pages)
                logger.info(f"Parsing PDF with {page_count} pages: {filename}")

                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text
                    page_text = page.extract_text() or ''
                    text_parts.append(page_text)
                    total_chars += len(page_text)

                    # Extract tables and convert to text
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            table_text = self._table_to_text(table)
                            text_parts.append(f"\n[Table]\n{table_text}\n")
                            total_chars += len(table_text) + 10

                    # Detect sections from page text
                    page_sections = self._detect_sections(page_text, page_num)

                    for section in page_sections:
                        if current_section:
                            sections.append(current_section)
                        current_section = section

                    # Clear page reference to help GC (pdfplumber caches page data)
                    page.flush_cache()

                    # Log progress for large PDFs
                    if page_count > 20 and page_num % 20 == 0:
                        logger.debug(f"Parsed {page_num}/{page_count} pages, {total_chars} chars")

            # Add final section
            if current_section:
                sections.append(current_section)

            # Join text parts into raw_text
            raw_text = '\n\n'.join(text_parts)

            # Clear text_parts list immediately after joining
            text_parts.clear()
            del text_parts

            # If no sections detected, create one from all content
            if not sections:
                sections = [{
                    'title': 'Document Content',
                    'content': raw_text,
                    'type': 'content',
                    'page': 1
                }]

            result = {
                'success': True,
                'raw_text': raw_text,
                'sections': sections,
                'metadata': {
                    'source_type': 'pdf',
                    'filename': filename,
                    'page_count': page_count,
                    'character_count': len(raw_text),
                    'section_count': len(sections)
                }
            }

            logger.info(f"PDF parsed: {page_count} pages, {len(raw_text)} chars, {len(sections)} sections")
            return result

        except Exception as e:
            logger.error(f"Error parsing PDF: {str(e)}")
            return {
                'success': False,
                'error': f'Failed to parse PDF: {str(e)}',
                'raw_text': '',
                'sections': [],
                'metadata': {}
            }
        finally:
            # Explicitly close and release BytesIO buffer
            if pdf_buffer is not None:
                pdf_buffer.close()
                del pdf_buffer

            # Force garbage collection after large PDF processing
            gc.collect()
            logger.debug("PDF parser memory cleanup complete")

    def parse_docx(self, file_content: bytes, filename: Optional[str] = None) -> Dict[str, Any]:
        """
        Parse Word document using python-docx.

        Extracts:
        - All paragraphs with style information
        - Headings (from paragraph styles)
        - Tables (converted to text)
        - Lists

        Args:
            file_content: Raw bytes of the DOCX file
            filename: Original filename

        Returns:
            Structured content dict
        """
        try:
            from docx import Document
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            return {
                'success': False,
                'error': 'Word document parsing library not available',
                'raw_text': '',
                'sections': [],
                'metadata': {}
            }

        try:
            doc = Document(BytesIO(file_content))

            all_text = []
            sections = []
            current_section = None
            current_content = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                all_text.append(text)

                # Check if this is a heading
                style_name = para.style.name.lower() if para.style else ''
                is_heading = 'heading' in style_name or 'title' in style_name

                # Also check for formatting-based headings
                if not is_heading:
                    is_heading = self._is_heading_text(text)

                if is_heading:
                    # Save previous section
                    if current_section:
                        current_section['content'] = '\n'.join(current_content)
                        sections.append(current_section)
                        current_content = []

                    # Determine heading level
                    level = 1
                    if 'heading' in style_name:
                        # Extract number from style name (e.g., "Heading 2" -> 2)
                        level_match = re.search(r'\d+', style_name)
                        if level_match:
                            level = int(level_match.group())

                    current_section = {
                        'title': text,
                        'type': 'heading',
                        'level': level,
                        'content': ''
                    }
                else:
                    current_content.append(text)

            # Add final section
            if current_section:
                current_section['content'] = '\n'.join(current_content)
                sections.append(current_section)
            elif current_content:
                # No headings found, create single section
                sections = [{
                    'title': 'Document Content',
                    'type': 'content',
                    'level': 1,
                    'content': '\n'.join(current_content)
                }]

            # Extract tables
            for table in doc.tables:
                table_text = self._docx_table_to_text(table)
                all_text.append(f"\n[Table]\n{table_text}\n")

            raw_text = '\n\n'.join(all_text)

            return {
                'success': True,
                'raw_text': raw_text,
                'sections': sections,
                'metadata': {
                    'source_type': 'docx',
                    'filename': filename,
                    'paragraph_count': len(doc.paragraphs),
                    'table_count': len(doc.tables),
                    'character_count': len(raw_text),
                    'section_count': len(sections)
                }
            }

        except Exception as e:
            logger.error(f"Error parsing DOCX: {str(e)}")
            return {
                'success': False,
                'error': f'Failed to parse Word document: {str(e)}',
                'raw_text': '',
                'sections': [],
                'metadata': {}
            }

    def parse_text(self, text: str, filename: Optional[str] = None) -> Dict[str, Any]:
        """
        Parse plain text content and detect structure.

        Uses pattern matching to identify:
        - Headings (markdown, numbered, ALL CAPS)
        - Sections/modules
        - Lessons/assignments

        Args:
            text: Plain text content
            filename: Original filename (optional)

        Returns:
            Structured content dict
        """
        if not text or not text.strip():
            return {
                'success': False,
                'error': 'Empty text content',
                'raw_text': '',
                'sections': [],
                'metadata': {}
            }

        lines = text.split('\n')
        sections = []
        current_section = None
        current_content = []

        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                current_content.append('')
                continue

            # Check if this line is a heading
            heading_info = self._detect_heading(line_stripped)

            if heading_info:
                # Save previous section
                if current_section:
                    current_section['content'] = '\n'.join(current_content).strip()
                    if current_section['content'] or current_section['title']:
                        sections.append(current_section)
                    current_content = []

                current_section = {
                    'title': heading_info['title'],
                    'type': heading_info['type'],
                    'level': heading_info.get('level', 1),
                    'content': ''
                }
            else:
                current_content.append(line_stripped)

        # Add final section
        if current_section:
            current_section['content'] = '\n'.join(current_content).strip()
            if current_section['content'] or current_section['title']:
                sections.append(current_section)
        elif current_content:
            # No structure detected, create single section
            sections = [{
                'title': 'Content',
                'type': 'content',
                'level': 1,
                'content': '\n'.join(current_content).strip()
            }]

        return {
            'success': True,
            'raw_text': text,
            'sections': sections,
            'metadata': {
                'source_type': 'text',
                'filename': filename,
                'line_count': len(lines),
                'character_count': len(text),
                'section_count': len(sections)
            }
        }

    def _detect_heading(self, text: str) -> Optional[Dict]:
        """
        Detect if a line is a heading and extract info.

        Returns:
            Dict with 'title', 'type', 'level' if heading detected, None otherwise
        """
        # Markdown headings: # Title, ## Title, etc.
        md_match = re.match(r'^(#{1,6})\s+(.+)$', text)
        if md_match:
            return {
                'title': md_match.group(2).strip(),
                'type': 'heading',
                'level': len(md_match.group(1))
            }

        # Numbered sections: 1. Title, 1.1 Title
        num_match = re.match(r'^(\d+(?:\.\d+)?)[.\)]\s+(.+)$', text)
        if num_match:
            level = 1 + text.count('.')
            return {
                'title': num_match.group(2).strip(),
                'type': 'numbered',
                'level': min(level, 3)
            }

        # Module/Unit/Chapter headers
        module_match = re.match(
            r'^(Module|Unit|Chapter|Lesson|Week|Section|Part)\s*(\d+)?[:\-\s]*(.*)$',
            text,
            re.IGNORECASE
        )
        if module_match:
            title_parts = [module_match.group(1)]
            if module_match.group(2):
                title_parts.append(module_match.group(2))
            if module_match.group(3):
                title_parts.append(f": {module_match.group(3).strip()}")
            return {
                'title': ''.join(title_parts),
                'type': 'module',
                'level': 1
            }

        # ALL CAPS headers (at least 3 words, mostly caps)
        if len(text) > 5 and len(text) < 100:
            upper_ratio = sum(1 for c in text if c.isupper()) / max(len(text.replace(' ', '')), 1)
            if upper_ratio > 0.7 and ' ' in text:
                return {
                    'title': text.title(),  # Convert to title case
                    'type': 'heading',
                    'level': 1
                }

        return None

    def _is_heading_text(self, text: str) -> bool:
        """Check if text appears to be a heading based on content patterns."""
        return self._detect_heading(text) is not None

    def _detect_sections(self, text: str, page_num: int) -> List[Dict]:
        """
        Detect section boundaries in a page of text.

        Args:
            text: Page text content
            page_num: Page number for reference

        Returns:
            List of section dicts
        """
        sections = []
        lines = text.split('\n')
        current_section = None
        current_content = []

        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                current_content.append('')
                continue

            heading_info = self._detect_heading(line_stripped)

            if heading_info:
                # Save previous section
                if current_section and (current_section['content'] or current_content):
                    current_section['content'] = '\n'.join(current_content).strip()
                    sections.append(current_section)
                    current_content = []

                current_section = {
                    'title': heading_info['title'],
                    'type': heading_info['type'],
                    'level': heading_info.get('level', 1),
                    'page': page_num,
                    'content': ''
                }
            else:
                current_content.append(line_stripped)

        # Don't forget the last section
        if current_section:
            current_section['content'] = '\n'.join(current_content).strip()
            sections.append(current_section)

        return sections

    def _table_to_text(self, table: List[List]) -> str:
        """
        Convert a pdfplumber table to readable text.

        Args:
            table: 2D list from pdfplumber table extraction

        Returns:
            Formatted text representation
        """
        if not table:
            return ''

        rows = []
        for row in table:
            cells = [str(cell or '').strip() for cell in row]
            rows.append(' | '.join(cells))

        return '\n'.join(rows)

    def _docx_table_to_text(self, table) -> str:
        """
        Convert a python-docx table to readable text.

        Args:
            table: Table object from python-docx

        Returns:
            Formatted text representation
        """
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(' | '.join(cells))

        return '\n'.join(rows)

    def detect_curriculum_type(self, sections: List[Dict]) -> str:
        """
        Analyze sections to determine what type of curriculum this likely is.

        Returns:
            String indicating type: 'syllabus', 'lesson_plan', 'course_outline',
            'assignment', 'textbook', 'unknown'
        """
        if not sections:
            return 'unknown'

        # Collect all text for analysis
        all_titles = ' '.join(s.get('title', '') for s in sections).lower()
        all_content = ' '.join(s.get('content', '') for s in sections).lower()
        all_text = all_titles + ' ' + all_content

        # Syllabus indicators
        syllabus_terms = ['syllabus', 'course objectives', 'grading policy', 'office hours',
                         'attendance', 'required materials', 'course schedule', 'prerequisites']
        syllabus_score = sum(1 for term in syllabus_terms if term in all_text)

        # Lesson plan indicators
        lesson_terms = ['learning objectives', 'lesson plan', 'warm-up', 'closure',
                       'materials needed', 'assessment', 'differentiation', 'time:']
        lesson_score = sum(1 for term in lesson_terms if term in all_text)

        # Course outline indicators
        outline_terms = ['module', 'unit', 'week', 'chapter', 'topics covered', 'reading']
        outline_score = sum(1 for term in outline_terms if term in all_text)

        # Assignment indicators
        assignment_terms = ['due date', 'submission', 'rubric', 'points', 'assignment',
                          'submit', 'deadline', 'requirements']
        assignment_score = sum(1 for term in assignment_terms if term in all_text)

        # Determine type
        scores = {
            'syllabus': syllabus_score,
            'lesson_plan': lesson_score,
            'course_outline': outline_score,
            'assignment': assignment_score
        }

        max_type = max(scores, key=scores.get)
        if scores[max_type] >= 2:
            return max_type

        return 'unknown'
